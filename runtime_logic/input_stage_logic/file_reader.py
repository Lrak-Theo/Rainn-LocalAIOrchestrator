import csv 
from PyPDF2 import PdfReader
from docx import Document

class FileReader:

    @staticmethod
    # read_file is the main function to determine which read method to use
    def read_file(file_path):
        file_path_lower = file_path.lower()

        if file_path_lower.endswith(".pdf"):
            return FileReader.read_pdf(file_path)
        elif file_path_lower.endswith(".txt"):
            return FileReader.read_txt(file_path)
        elif file_path_lower.endswith(".csv"):
            return FileReader.read_csv(file_path)
        elif file_path_lower.endswith(".docx"):
            return FileReader.read_docx(file_path)
        else:
            return "[Unsupported file type]"

    @staticmethod
    #This is how using PyPDF2 can enable Rainn to read pdf files
    #Basically a text extractor which is added into a var (or array) combined with breaks using \n
    def read_pdf(path):
        # User the library to read the pdf
        reader = PdfReader(path)
        
        # Begin to convert the text to plain text
        plain_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                # Combine all string with new lines
                plain_text += text + '\n'
        
        plain_text = plain_text.replace("\r\n", "\n").replace("\r", "\n")

        if "\n" not in plain_text and len(plain_text) > 250:
            plain_text = plain_text.replace(". ", ".\n")
        
        # Return the now converted string
        return plain_text

    # Simple helper for reading txt files from the project directory.
    @staticmethod
    def read_txt(path):
        # Txt files requires no extra operations to read the file
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    # CSV reader function
    @staticmethod
    def read_csv(path):
        # Create an empty list
        rows = []

        # For row by row, read the line and append the line to the list
        with open(path, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))

        return "\n".join(rows)

    # Ms Word reader function
    @staticmethod
    def read_docx(path):
        # Read the docx file and have it assigned as an object
        doc = Document(path)

        # Empty list to being line by line conversion
        chunks = []

        # Read paragraph by paragraph
        for paragraph in doc.paragraphs:

            # Remove the leading and trailing whitesapce and then add it to the list
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        # Tables in docx files will need to be converted 
        # Get a table in the docx
        for table in doc.tables:

            # Read row by row of that table
            for row in table.rows:
                # Create a cells list
                cells = []

                # Then in each cell in the row
                for cell in row.cells:
                    # Split cell text into lines -> trim each line -> drop empty lines
                    cell_text = " ".join(
                        line.strip()
                        for line in cell.text.splitlines()
                        if line.strip()
                    )
                    cells.append(cell_text)

                # If at least one cell has content, append a row string to chunks as comma-seperated cells.
                if any(cells):
                    chunks.append(", ".join(cells))

        # Return the converted docx file
        return "\n".join(chunks).strip()
